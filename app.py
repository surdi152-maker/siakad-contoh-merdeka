import streamlit as st
import pandas as pd
import sqlite3
import bcrypt
import os
import io
from datetime import date, datetime
from fpdf import FPDF
from docx import Document

# ==========================================
# 1. KONFIGURASI UTAMA
# ==========================================

st.set_page_config(
    page_title="SIAKAD SMPN 152", 
    page_icon="🏫", 
    layout="wide",
    initial_sidebar_state="expanded"  # Ini kuncinya agar menu selalu muncul
)

def set_custom_css():
    st.markdown("""
        <style>
        /* 1. Sembunyikan elemen bawaan Streamlit yang tidak perlu */
        #MainMenu {visibility: hidden;}
        footer {visibility: hidden;}
        .stDeployButton {display:none;}
        
        /* 2. Sembunyikan tombol 'Manage App' / Status Widget (KUNCI UTAMA) */
        [data-testid="stStatusWidget"] {
            visibility: hidden !important;
            display: none !important;
        }
        
        /* 3. Desain tombol UI Anda */
        div.stButton > button:first-child {
            background-color: #D4AF37; color: #0F1115; border-radius: 8px; border: none;
            box-shadow: 0 4px 15px rgba(212, 175, 55, 0.2); transition: all 0.3s ease;
            font-weight: 700; padding: 0.6rem 1.2rem;
        }
        div.stButton > button:first-child:hover {
            background-color: #AA841F; color: white;
            box-shadow: 0 8px 20px rgba(212, 175, 55, 0.4); transform: translateY(-2px);
        }
        [data-testid="stSidebar"] { box-shadow: 4px 0 15px rgba(0,0,0,0.5); border-right: 1px solid #2A2E39; }
        </style>
    """, unsafe_allow_html=True)

set_custom_css()

DB_FILE = 'sekolah_lokal.db'
DAFTAR_KELAS = [
    "7-A", "7-B", "8-A", "8-B", "9-A", "9-B",
]

TARIF_SPP_BULANAN = 250000
DAFTAR_BULAN_SPP = [
    "Juli", "Agustus", "September", "Oktober", "November", "Desember",
    "Januari", "Februari", "Maret", "April", "Mei", "Juni"
]

# ==========================================
# 2. MESIN EXPORT FILE (EXCEL, WORD, PDF)
# ==========================================

def bersihkan_teks_untuk_pdf(teks):
    return str(teks).encode('latin-1', 'ignore').decode('latin-1')

def export_excel(df):
    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine='openpyxl') as writer:
        df.to_excel(writer, index=False)
    return buf.getvalue()

def export_word(df, judul):
    doc = Document()
    doc.add_heading(judul, level=1)
    tabel = doc.add_table(rows=1, cols=len(df.columns))
    tabel.style = 'Table Grid'
    hdr_cells = tabel.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col).upper()
    for _, row in df.iterrows():
        row_cells = tabel.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

def export_pdf(df, judul):
    pdf = FPDF(orientation='L', unit='mm', format='A4')
    pdf.add_page()
    pdf.set_font("Arial", 'B', 12)
    
    judul_bersih = bersihkan_teks_untuk_pdf(judul)
    pdf.cell(0, 10, judul_bersih, ln=True, align='C')
    pdf.ln(5)
    pdf.set_font("Arial", 'B', 8)
    
    lebar_halaman = 277 
    lebar_kolom = lebar_halaman / len(df.columns) if not df.empty else 20
    
    for col in df.columns:
        kolom_bersih = bersihkan_teks_untuk_pdf(str(col).upper())
        pdf.cell(lebar_kolom, 8, kolom_bersih, border=1, align='C')
    pdf.ln()
    
    pdf.set_font("Arial", '', 8)
    for _, row in df.iterrows():
        for val in row:
            val_bersih = bersihkan_teks_untuk_pdf(str(val)[:25])
            pdf.cell(lebar_kolom, 8, val_bersih, border=1, align='L')
        pdf.ln()
    return pdf.output(dest="S").encode("latin-1", 'replace')

def tampilkan_tombol_download(df, nama_file, judul_dokumen):
    st.markdown("⬇️ **Download Tabel Ini:**")
    k1, k2, k3, _ = st.columns([1.5, 1.5, 1.5, 5])
    with k1:
        st.download_button("📊 Excel", export_excel(df), file_name=f"{nama_file}.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", key=f"ex_{nama_file}")
    with k2:
        st.download_button("📝 Word", export_word(df, judul_dokumen), file_name=f"{nama_file}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key=f"wd_{nama_file}")
    with k3:
        st.download_button("📄 PDF", export_pdf(df, judul_dokumen), file_name=f"{nama_file}.pdf", mime="application/pdf", key=f"pf_{nama_file}")
    st.markdown("<br>", unsafe_allow_html=True)

def dapatkan_deskripsi_kompetensi(mapel, nilai):
    if nilai >= 85:
        return f"Menunjukkan penguasaan sangat baik dalam memahami materi {mapel} serta mampu menerapkannya secara mandiri."
    elif nilai >= 70:
        return f"Menunjukkan penguasaan baik dalam kompetensi materi {mapel} dan mampu mengikuti pembelajaran dengan lancar."
    else:
        return f"Perlu bimbingan intensif dan peningkatan fokus terutama dalam memahami konsep dasar pada materi {mapel}."

# ==========================================
# 3. DATABASE & AUTENTIKASI
# ==========================================
def get_connection():
    return sqlite3.connect(DB_FILE)

def init_db():
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute('''CREATE TABLE IF NOT EXISTS users (id INTEGER PRIMARY KEY AUTOINCREMENT, username TEXT UNIQUE NOT NULL, password_hash TEXT NOT NULL, role TEXT NOT NULL)''')
    cursor.execute('''CREATE TABLE IF NOT EXISTS absensi_siswa (id INTEGER PRIMARY KEY AUTOINCREMENT, tanggal TEXT NOT NULL, nis TEXT, nama_siswa TEXT NOT NULL, kelas TEXT NOT NULL, status_kehadiran TEXT NOT NULL, petugas TEXT NOT NULL)''')
    cursor.execute('''CREATE TABLE IF NOT EXISTS absensi_guru (id INTEGER PRIMARY KEY AUTOINCREMENT, tanggal TEXT NOT NULL, username_guru TEXT NOT NULL, waktu_masuk TEXT, waktu_pulang TEXT, status_kehadiran TEXT NOT NULL)''')
    cursor.execute('''CREATE TABLE IF NOT EXISTS buku_nilai (id INTEGER PRIMARY KEY AUTOINCREMENT, tanggal_input TEXT NOT NULL, kelas TEXT NOT NULL, mata_pelajaran TEXT NOT NULL, nis TEXT, nama_siswa TEXT NOT NULL, nilai_absensi REAL, nilai_tugas REAL, nilai_uts REAL, nilai_uas REAL, nilai_akhir REAL, guru_penginput TEXT NOT NULL)''')
    cursor.execute('''CREATE TABLE IF NOT EXISTS pembayaran_spp (id INTEGER PRIMARY KEY AUTOINCREMENT, tanggal_bayar TEXT NOT NULL, nama_siswa TEXT NOT NULL, kelas TEXT NOT NULL, bulan TEXT NOT NULL, jumlah_bayar REAL NOT NULL, status TEXT NOT NULL, petugas TEXT NOT NULL)''')
    
    cursor.execute("SELECT COUNT(*) FROM users WHERE username = 'admin'")
    if cursor.fetchone()[0] == 0:
        hashed_default = bcrypt.hashpw("admin123".encode('utf-8'), bcrypt.gensalt()).decode('utf-8')
        cursor.execute("INSERT INTO users (username, password_hash, role) VALUES ('admin', ?, 'admin')", (hashed_default,))
    conn.commit()
    conn.close()

init_db()

def check_password(password, hashed):
    return bcrypt.checkpw(password.encode('utf-8'), hashed.encode('utf-8'))

if 'logged_in' not in st.session_state:
    st.session_state['logged_in'] = False
    st.session_state['username'] = ''
    st.session_state['role'] = ''

def logout_user():
    st.session_state['logged_in'] = False
    st.session_state['username'] = ''
    st.session_state['role'] = ''
    st.rerun()

# ==========================================
# 4. GERBANG LOGIN
# ==========================================
if not st.session_state['logged_in']:
    st.title("🔐 Login SIAKAD SMPN 152")
       
    with st.form("login_form"):
        user_input = st.text_input("Username")
        pass_input = st.text_input("Password", type="password")
        if st.form_submit_button("Masuk Ke Sistem"):
            user_input_bersih = user_input.strip().lower()
            conn = get_connection()
            cursor = conn.cursor()
            cursor.execute("SELECT password_hash, role FROM users WHERE username = ?", (user_input_bersih,))
            user_data = cursor.fetchone()
            conn.close()
            
            if user_data and check_password(pass_input, user_data[0]):
                st.session_state['logged_in'] = True
                st.session_state['username'] = user_input_bersih
                st.session_state['role'] = user_data[1]
                st.rerun()
            else:
                st.error("❌ Username atau password salah.")

# ==========================================
# 5. HALAMAN UTAMA APLIKASI
# ==========================================
else:
    st.title("🏫 SIAKAD & TU")
    
    st.sidebar.title("Navigasi")
    st.sidebar.success(f"👤 Akun: **{st.session_state['username']}** ({st.session_state['role'].upper()})")
    if st.sidebar.button("Keluar (Logout)"):
        logout_user()
    st.sidebar.markdown("---")
    
    menu = ["Absensi Guru", "Absensi Siswa", "Lihat Laporan Guru dan Siswa", "Buku Nilai (Upload)", "Laporan & E-Rapor"]
    if st.session_state['role'] == 'admin':
        menu.extend(["Manajemen SPP", "Manajemen Akun Guru", "Koreksi & Hapus Data", "Backup & Restore Data"])
        
    pilihan = st.sidebar.radio("Pilih Menu Modul:", menu)

    # --- MODUL 1: ABSENSI GURU ---
    if pilihan == "Absensi Guru":
        st.header("👨‍🏫 Portal Kehadiran Guru")
        hari_ini = str(date.today())
        waktu_sekarang = datetime.now().strftime("%H:%M:%S")
        username_aktif = st.session_state['username']
        
        conn = get_connection()
        cursor = conn.cursor()
        cursor.execute("SELECT waktu_masuk, waktu_pulang FROM absensi_guru WHERE tanggal = ? AND username_guru = ?", (hari_ini, username_aktif))
        data_absen = cursor.fetchone()
        
        if not data_absen:
            st.info("Anda belum melakukan absensi masuk hari ini.")
            with st.form("form_absen_masuk"):
                status_guru = st.radio("Status Hari Ini", ["Hadir", "Izin", "Sakit"], horizontal=True)
                if st.form_submit_button("Catat Jam Masuk"):
                    jam_masuk = waktu_sekarang if status_guru == "Hadir" else "-"
                    cursor.execute('INSERT INTO absensi_guru (tanggal, username_guru, waktu_masuk, waktu_pulang, status_kehadiran) VALUES (?, ?, ?, ?, ?)', 
                                   (hari_ini, username_aktif, jam_masuk, "-", status_guru))
                    conn.commit()
                    st.success("Absen masuk berhasil disimpan.")
                    st.rerun()
        elif data_absen[0] != "-" and data_absen[1] == "-":
            st.success(f"✅ Anda sudah absen masuk pada jam **{data_absen[0]}**.")
            with st.form("form_absen_pulang"):
                if st.form_submit_button("Catat Jam Pulang"):
                    cursor.execute('UPDATE absensi_guru SET waktu_pulang = ? WHERE tanggal = ? AND username_guru = ?', (waktu_sekarang, hari_ini, username_aktif))
                    conn.commit()
                    st.success("Absen pulang berhasil disimpan.")
                    st.rerun()
        else:
            st.success("🎉 Anda telah menyelesaikan absensi untuk hari ini.")
        conn.close()

    # --- MODUL 2: ABSENSI SISWA ---
    elif pilihan == "Absensi Siswa":
        st.header("📝 Formulir Kehadiran Siswa")
        
        conn = get_connection()
        df_siswa_lama = pd.read_sql_query("SELECT DISTINCT nis, nama_siswa, kelas FROM absensi_siswa ORDER BY nama_siswa ASC", conn)
        conn.close()
        
        pilihan_cepat = ["-- Input Siswa Baru / Ketik Manual --"]
        if not df_siswa_lama.empty:
            df_siswa_lama['display'] = df_siswa_lama['nis'].astype(str) + " - " + df_siswa_lama['nama_siswa'] + " (" + df_siswa_lama['kelas'] + ")"
            pilihan_cepat.extend(df_siswa_lama['display'].tolist())
            
        siswa_terpilih = st.selectbox("🎯 Pilih dari Siswa yang Sudah Pernah Absen (Cepat):", pilihan_cepat)
        
        nama_auto = ""
        kelas_auto = DAFTAR_KELAS[0]
        nis_auto = ""
        
        if siswa_terpilih != "-- Input Siswa Baru / Ketik Manual --":
            nis_auto = siswa_terpilih.split(" - ")[0]
            row_siswa = df_siswa_lama[df_siswa_lama['nis'] == nis_auto].iloc[0]
            nama_auto = row_siswa['nama_siswa']
            kelas_auto = row_siswa['kelas']
            st.success(f"✅ Memuat data otomatis untuk: **{nama_auto}**")
        
        nis_input = st.text_input("Masukkan NIS Siswa", value=nis_auto)
        
        if nis_input and siswa_terpilih == "-- Input Siswa Baru / Ketik Manual --":
            conn = get_connection()
            cursor = conn.cursor()
            cursor.execute("SELECT nama_siswa, kelas FROM absensi_siswa WHERE nis = ? LIMIT 1", (nis_input.strip(),))
            data_siswa = cursor.fetchone()
            conn.close()
            
            if data_siswa:
                nama_auto = data_siswa[0]
                kelas_auto = data_siswa[1]
                st.success(f"✅ NIS Terdeteksi: **{nama_auto}**")
        
        with st.form("form_absensi_manual", clear_on_submit=True):
            tgl_absen = st.date_input("Tanggal Kehadiran", date.today())
            nm_input = st.text_input("Nama Lengkap Siswa", value=nama_auto)
            idx_kelas = DAFTAR_KELAS.index(kelas_auto) if kelas_auto in DAFTAR_KELAS else 0
            kls_input = st.selectbox("Pilih Kelas", DAFTAR_KELAS, index=idx_kelas)
            st_kehadiran = st.radio("Status Kehadiran:", ["Hadir", "Sakit", "Izin", "Alpa"], horizontal=True)
            
            if st.form_submit_button("Simpan Absensi"):
                if not nis_input.strip() or not nm_input.strip():
                    st.error("❌ NIS dan Nama Siswa wajib diisi!")
                else:
                    conn = get_connection()
                    cursor = conn.cursor()
                    cursor.execute('''INSERT INTO absensi_siswa (tanggal, nis, nama_siswa, kelas, status_kehadiran, petugas) 
                                      VALUES (?, ?, ?, ?, ?, ?)''',
                                   (str(tgl_absen), nis_input.strip(), nm_input.strip(), kls_input, st_kehadiran, st.session_state['username']))
                    conn.commit()
                    conn.close()
                    st.success(f"✅ Absensi {nm_input} ({nis_input}) berhasil disimpan!")
                    st.rerun()

    # --- MODUL 3: DASHBOARD MONITORING ---
    elif pilihan == "Lihat Laporan Guru dan Siswa":
        st.header("📊 Dashboard Monitoring & Laporan Kehadiran")
        conn = get_connection()
        df_siswa = pd.read_sql_query("SELECT tanggal, nis, nama_siswa, kelas, status_kehadiran, petugas FROM absensi_siswa", conn)
        df_guru = pd.read_sql_query("SELECT tanggal, username_guru, waktu_masuk, waktu_pulang, status_kehadiran FROM absensi_guru", conn)
        conn.close()
        
        tab_guru, tab_siswa = st.tabs(["👨‍🏫 Laporan Absensi Guru", "👶 Laporan Absensi Siswa"])
        
        with tab_guru:
            if df_guru.empty:
                st.info("Belum ada data kehadiran guru.")
            else:
                col1, col2 = st.columns([1, 2])
                with col1:
                    st.markdown("### 📈 Ringkasan")
                    st.metric("Total Input Absen", f"{len(df_guru)} Data")
                    st.metric("Guru Hadir", f"{len(df_guru[df_guru['status_kehadiran'] == 'Hadir'])} Orang")
                with col2:
                    st.markdown("### 📊 Grafik")
                    st.bar_chart(df_guru['status_kehadiran'].value_counts())
                st.markdown("---")
                st.dataframe(df_guru, use_container_width=True)
                tampilkan_tombol_download(df_guru, "Laporan_Absen_Guru", "REKAPITULASI ABSENSI GURU")

        with tab_siswa:
            if df_siswa.empty:
                st.info("Belum ada data kehadiran siswa.")
            else:
                col1, col2 = st.columns([1, 2])
                with col1:
                    st.markdown("### 📈 Ringkasan")
                    st.metric("Total Input Absen", f"{len(df_siswa)} Data")
                    st.metric("Siswa Hadir", f"{len(df_siswa[df_siswa['status_kehadiran'] == 'Hadir'])} Orang")
                with col2:
                    st.markdown("### 📊 Grafik")
                    st.bar_chart(df_siswa['status_kehadiran'].value_counts())
                st.markdown("---")
                st.dataframe(df_siswa, use_container_width=True)
                tampilkan_tombol_download(df_siswa, "Laporan_Absen_Siswa", "REKAPITULASI ABSENSI SISWA")

    # --- MODUL 4: BUKU NILAI (UPLOAD) ---
    elif pilihan == "Buku Nilai (Upload)":
        st.header("📈 Upload Nilai Siswa (Otomatis)")
        df_template = pd.DataFrame({"NIS": ["1001", "1002"], "Nama Siswa": ["Suntoro", "Lestari"], "Nilai Absensi": [100, 95], "Nilai Tugas": [85, 78], "Nilai UTS": [80, 85], "Nilai UAS": [90, 88]})
        buf = io.BytesIO()
        with pd.ExcelWriter(buf, engine='openpyxl') as writer:
            df_template.to_excel(writer, index=False, sheet_name='Nilai')
        st.download_button(label="📥 Download Template Excel", data=buf.getvalue(), file_name="Template_Nilai_NIS.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        
        st.markdown("---")
        with st.form("form_upload_nilai"):
            kls_pilih = st.selectbox("Pilih Kelas", DAFTAR_KELAS)
            mapel_pilih = st.text_input("Mata Pelajaran (Contoh: Matematika)")
            file_excel = st.file_uploader("Unggah Berkas Excel (.xlsx)", type=["xlsx"])
            
            if st.form_submit_button("Proses & Ekstraksi"):
                if file_excel and mapel_pilih.strip():
                    try:
                        df_in = pd.read_excel(file_excel)
                        df_in['Nilai Akhir'] = (df_in['Nilai Absensi'] * 0.1) + (df_in['Nilai Tugas'] * 0.2) + (df_in['Nilai UTS'] * 0.3) + (df_in['Nilai UAS'] * 0.4)
                        
                        for col in ['Nilai Absensi', 'Nilai Tugas', 'Nilai UTS', 'Nilai UAS', 'Nilai Akhir']:
                            if col in df_in.columns:
                                df_in[col] = df_in[col].round().astype(int)

                        conn = get_connection()
                        cursor = conn.cursor()
                        tgl_skrg = str(date.today())
                        for _, row in df_in.iterrows():
                            cursor.execute('''INSERT INTO buku_nilai (tanggal_input, kelas, mata_pelajaran, nis, nama_siswa, nilai_absensi, nilai_tugas, nilai_uts, nilai_uas, nilai_akhir, guru_penginput)
                                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)''', 
                                (tgl_skrg, kls_pilih, mapel_pilih.upper(), str(row['NIS']), str(row['Nama Siswa']), float(row['Nilai Absensi']), float(row['Nilai Tugas']), float(row['Nilai UTS']), float(row['Nilai UAS']), float(row['Nilai Akhir']), st.session_state['username']))
                        conn.commit()
                        conn.close()
                        st.success(f"Berhasil menyimpan {len(df_in)} data nilai siswa.")
                        st.dataframe(df_in, use_container_width=True)
                    except Exception as e:
                        st.error(f"Error pembacaan berkas: {e}")

    # --- MODUL 5: E-RAPOR ---
    elif pilihan == "Laporan & E-Rapor":
        st.header("🖨️ Cetak Hasil Belajar Siswa (E-Rapor)")
        c_kelas = st.selectbox("Pilih Kelas", DAFTAR_KELAS)
        
        conn = get_connection()
        df_siswa_kelas = pd.read_sql_query("SELECT DISTINCT nis, nama_siswa FROM buku_nilai WHERE kelas = ?", conn, params=(c_kelas,))
        conn.close()
        
        if df_siswa_kelas.empty:
            st.info(f"Belum ada data nilai untuk kelas {c_kelas}.")
        else:
            daftar_pilihan = df_siswa_kelas['nis'].astype(str) + " - " + df_siswa_kelas['nama_siswa']
            pilihan_siswa = st.selectbox("Pilih Siswa (NIS & Nama)", options=daftar_pilihan)
            
            if st.button("Cari & Tampilkan Laporan Rapor"):
                nis_terpilih = pilihan_siswa.split(" - ")[0]
                
                conn = get_connection()
                df_rp = pd.read_sql_query("SELECT mata_pelajaran, nilai_akhir FROM buku_nilai WHERE kelas = ? AND nis = ?", conn, params=(c_kelas, nis_terpilih))
                
                cursor = conn.cursor()
                cursor.execute("SELECT status_kehadiran, COUNT(*) FROM absensi_siswa WHERE nis = ? GROUP BY status_kehadiran", (nis_terpilih,))
                rekap_absen = dict(cursor.fetchall())
                conn.close()
                
                sakit = rekap_absen.get('Sakit', 0)
                izin = rekap_absen.get('Izin', 0)
                alpa = rekap_absen.get('Alpa', 0)
                
                if df_rp.empty:
                    st.warning("Data nilai siswa tidak ditemukan.")
                else:
                    df_rp['nilai_akhir'] = df_rp['nilai_akhir'].round().astype(int)
                    df_rp['Capaian Kompetensi'] = df_rp.apply(lambda r: dapatkan_deskripsi_kompetensi(r['mata_pelajaran'], r['nilai_akhir']), axis=1)
                    st.dataframe(df_rp, use_container_width=True)
                    tampilkan_tombol_download(df_rp, f"Tabel_Rapor_{nis_terpilih}", f"PRATINJAU RAPOR SISWA: {pilihan_siswa}")
                    
                    st.markdown("---")
                    st.markdown("### 🖨️ Cetak Dokumen Rapor Resmi")
                    
                    nis_cetak = nis_terpilih
                    nama_cetak = pilihan_siswa.split(" - ")[1]
                    
                    pdf = FPDF(orientation='P', unit='mm', format='A4')
                    pdf.add_page()
                    
                    logo = "logo.png"
                    if os.path.exists(logo):
                        pdf.image(logo, x=12, y=10, w=22)
                        
                    pdf.set_y(10)
                    pdf.set_font("Arial", 'B', 14)
                    pdf.cell(0, 7, txt="YAYASAN PENDIDIKAN NASIONAL", ln=True, align='C')
                    pdf.set_font("Arial", 'B', 16)
                    pdf.cell(0, 7, txt="SMPN 152", ln=True, align='C')
                    pdf.set_font("Arial", 'I', 9)
                    pdf.cell(0, 5, txt="Jl. Raya Lembaga Sekolah No. 12, Jati Agung, Lampung Selatan, 35365 - Telp: (0721) 5554321", ln=True, align='C')
                    pdf.ln(5)
                    pdf.line(10, 36, 200, 36)
                    pdf.line(10, 37, 200, 37)
                    pdf.ln(8)
                    
                    pdf.set_font("Arial", 'B', 12)
                    pdf.cell(0, 10, txt="HASIL BELAJAR SISWA (E-RAPOR)", ln=True, align='C')
                    pdf.ln(4)
                    
                    pdf.set_font("Arial", size=10)
                    pdf.cell(35, 6, "Nama Peserta Didik", 0)
                    pdf.cell(5, 6, ": ", 0)
                    pdf.set_font("Arial", 'B', 10)
                    pdf.cell(0, 6, bersihkan_teks_untuk_pdf(nama_cetak).upper(), 0, 1)
                    
                    pdf.set_font("Arial", size=10)
                    pdf.cell(35, 6, "NIS", 0)
                    pdf.cell(5, 6, ": ", 0)
                    pdf.cell(0, 6, str(nis_cetak), 0, 1)
                    
                    pdf.cell(35, 6, "Kelas", 0)
                    pdf.cell(5, 6, ": ", 0)
                    pdf.cell(0, 6, c_kelas, 0, 1)
                    pdf.ln(5)
                    
                    pdf.set_font("Arial", 'B', 9)
                    pdf.cell(10, 8, "No", 1, 0, 'C')
                    pdf.cell(50, 8, "Mata Pelajaran", 1, 0, 'C')
                    pdf.cell(25, 8, "Nilai Akhir", 1, 0, 'C')
                    pdf.cell(105, 8, "Capaian Kompetensi (Deskripsi)", 1, 1, 'C')
                    
                    pdf.set_font("Arial", size=9)
                    no_urut = 1
                    for _, row in df_rp.iterrows():
                        nilai_int = int(row['nilai_akhir'])
                        desc_text = bersihkan_teks_untuk_pdf(dapatkan_deskripsi_kompetensi(row['mata_pelajaran'], nilai_int))
                        mapel_bersih = bersihkan_teks_untuk_pdf(row['mata_pelajaran'])
                        
                        x_pos = pdf.get_x()
                        y_pos = pdf.get_y()
                        
                        pdf.cell(10, 14, str(no_urut), border=1, align='C')
                        pdf.cell(50, 14, mapel_bersih, border=1, align='L')
                        pdf.cell(25, 14, str(nilai_int), border=1, align='C')
                        
                        pdf.set_xy(x_pos + 85, y_pos)
                        pdf.multi_cell(105, 4.6, desc_text, border=1, align='L')
                        
                        pdf.set_xy(x_pos, y_pos + 14)
                        no_urut += 1
                        
                    pdf.ln(6)
                    
                    pdf.set_font("Arial", 'B', 10)
                    pdf.cell(0, 6, "Ketidakhadiran (Absensi)", 0, 1, 'L')
                    pdf.set_font("Arial", size=9)
                    # PERUBAHAN: Lebar kolom dinaikkan menjadi 55 agar teks tidak terpotong
                    pdf.cell(55, 6, "1. Sakit", 1, 0, 'L')
                    pdf.cell(25, 6, f"{sakit} Hari", 1, 1, 'C')
                    pdf.cell(55, 6, "2. Izin", 1, 0, 'L')
                    pdf.cell(25, 6, f"{izin} Hari", 1, 1, 'C')
                    pdf.cell(55, 6, "3. Tanpa Keterangan (Alpa)", 1, 0, 'L')
                    pdf.cell(25, 6, f"{alpa} Hari", 1, 1, 'C')
                    
                    pdf_out = pdf.output(dest="S").encode("latin-1", 'replace')
                    st.download_button(label="📄 Cetak & Download PDF Lembar Rapor Resmi", data=pdf_out, file_name=f"Rapor_{nama_cetak.strip()}.pdf", mime="application/pdf")

    # --- MODUL 6: MANAJEMEN SPP ---
    elif pilihan == "Manajemen SPP" and st.session_state['role'] == 'admin':
        st.header("💰 Sistem Manajemen & Pemantauan Tunggakan SPP")
        
        tab_input, tab_buku_besar, tab_rekap_tunggakan = st.tabs([
            "📝 Input Kas Pembayaran", 
            "👤 Buku Besar Per Siswa", 
            "📋 Master Laporan Tunggakan Sekolah"
        ])
        
        with tab_input:
            st.subheader("Formulir Entri Pembayaran SPP Resmi")
            with st.form("form_spp_baru", clear_on_submit=True):
                tgl_sp = st.date_input("Tanggal Setor Tunai", date.today())
                nm_sp = st.text_input("Nama Lengkap Siswa")
                kls_sp = st.selectbox("Kelas Aktif", DAFTAR_KELAS)
                bln_sp = st.selectbox("Untuk Pembayaran Bulan", DAFTAR_BULAN_SPP)
                nom_sp = st.number_input("Nominal Pembayaran (Rp)", min_value=0, value=TARIF_SPP_BULANAN, step=10000)
                st_sp = st.radio("Status Verifikasi Berkas", ["LUNAS", "PENDING"], horizontal=True)
                
                if st.form_submit_button("Simpan & Cetak Kwitansi"):
                    if nm_sp.strip():
                        conn = get_connection()
                        cursor = conn.cursor()
                        cursor.execute('INSERT INTO pembayaran_spp (tanggal_bayar, nama_siswa, kelas, bulan, jumlah_bayar, status, petugas) VALUES (?, ?, ?, ?, ?, ?, ?)',
                                       (str(tgl_sp), nm_sp.strip().upper(), kls_sp, bln_sp, nom_sp, st_sp, st.session_state['username']))
                        conn.commit()
                        conn.close()
                        st.success(f"✅ Data Pembayaran SPP {nm_sp.upper()} bulan {bln_sp} sebesar Rp {nom_sp:,.0f} sukses diverifikasi!")
                    else:
                        st.error("❌ Nama siswa tidak boleh kosong!")

        with tab_buku_besar:
            st.subheader("Lembar Kendali & Riwayat Tunggakan Individu")
            conn = get_connection()
            df_all_siswa = pd.read_sql_query("SELECT DISTINCT nama_siswa FROM pembayaran_spp ORDER BY nama_siswa ASC", conn)
            conn.close()
            
            if df_all_siswa.empty:
                st.info("Belum ada mutasi keuangan siswa di database.")
            else:
                pilih_siswa_spp = st.selectbox("Cari Nama Siswa:", df_all_siswa['nama_siswa'].tolist())
                
                if pilih_siswa_spp:
                    conn = get_connection()
                    df_bayar_siswa = pd.read_sql_query("SELECT tanggal_bayar, kelas, bulan, jumlah_bayar, status, petugas FROM pembayaran_spp WHERE nama_siswa = ?", conn, params=(pilih_siswa_spp,))
                    conn.close()
                    
                    data_matriks_spp = []
                    total_tunggakan_siswa = 0
                    total_terbayar_siswa = 0
                    
                    for bulan in DAFTAR_BULAN_SPP:
                        kondisi_bulan = df_bayar_siswa[(df_bayar_siswa['bulan'] == bulan) & (df_bayar_siswa['status'] == 'LUNAS')]
                        if not kondisi_bulan.empty:
                            row = kondisi_bulan.iloc[0]
                            data_matriks_spp.append({
                                "Bulan": bulan,
                                "Status": "LUNAS",
                                "Tanggal Bayar": row['tanggal_bayar'],
                                "Jumlah Setor": f"Rp {row['jumlah_bayar']:,.0f}",
                                "Petugas TU": row['petugas']
                            })
                            total_terbayar_siswa += row['jumlah_bayar']
                        else:
                            data_matriks_spp.append({
                                "Bulan": bulan,
                                "Status": "TUNGGAKAN",
                                "Tanggal Bayar": "-",
                                "Jumlah Setor": f"Rp {TARIF_SPP_BULANAN:,.0f}",
                                "Petugas TU": "-"
                            })
                            total_tunggakan_siswa += TARIF_SPP_BULANAN
                            
                    df_matriks_final = pd.DataFrame(data_matriks_spp)
                    
                    c1, c2 = st.columns(2)
                    c1.metric("Total Dana Terbayar", f"Rp {total_terbayar_siswa:,.0f}")
                    c2.metric("Sisa Beban Tunggakan", f"Rp {total_tunggakan_siswa:,.0f}", delta=f"-Rp {total_tunggakan_siswa:,.0f}" if total_tunggakan_siswa > 0 else "0")
                    
                    st.markdown(f"### 📋 Kartu Kendali SPP: {pilih_siswa_spp}")
                    st.dataframe(df_matriks_final, use_container_width=True)
                    tampilkan_tombol_download(df_matriks_final, f"Buku_Besar_SPP_{pilih_siswa_spp}", f"KARTU KENDALI SPP INDIVIDU: {pilih_siswa_spp}")

        with tab_rekap_tunggakan:
            st.subheader("Daftar Buku Piutang & Rekapitulasi Tunggakan Sekolah")
            conn = get_connection()
            df_master_spp = pd.read_sql_query("SELECT nama_siswa, kelas, bulan, jumlah_bayar, status FROM pembayaran_spp", conn)
            conn.close()
            
            if df_master_spp.empty:
                st.info("Tidak ada data penunggak, semua administrasi bersih.")
            else:
                semua_siswa_unik = df_master_spp['nama_siswa'].unique()
                rekap_tunggakan_sekolah = []
                
                for siswa in semua_siswa_unik:
                    df_sub = df_master_spp[df_master_spp['nama_siswa'] == siswa]
                    kelas_siswa = df_sub['kelas'].iloc[0]
                    
                    bulan_menunggak = []
                    nominal_tunggakan = 0
                    
                    for bulan in DAFTAR_BULAN_SPP:
                        if df_sub[(df_sub['bulan'] == bulan) & (df_sub['status'] == 'LUNAS')].empty:
                            bulan_menunggak.append(bulan)
                            nominal_tunggakan += TARIF_SPP_BULANAN
                            
                    if nominal_tunggakan > 0:
                        rekap_tunggakan_sekolah.append({
                            "Nama Siswa": siswa,
                            "Kelas": kelas_siswa,
                            "Total Tunggakan (Rp)": nominal_tunggakan,
                            "Detail Bulan Menunggak": ", ".join(bulan_menunggak)
                        })
                        
                if rekap_tunggakan_sekolah:
                    df_rekap_tunggakan_final = pd.DataFrame(rekap_tunggakan_sekolah)
                    
                    grand_total_piutang = df_rekap_tunggakan_final['Total Tunggakan (Rp)'].sum()
                    st.error(f"⚠️ Total Piutang/Tunggakan SPP Belum Masuk Kas Sekolah: **Rp {grand_total_piutang:,.0f}**")
                    
                    st.dataframe(df_rekap_tunggakan_final, use_container_width=True)
                    tampilkan_tombol_download(df_rekap_tunggakan_final, "Master_Laporan_Tunggakan_SPP", "LAPORAN PIUTANG DAN TUNGGAKAN SPP SEKOLAH")
                else:
                    st.success("🎉 Luar biasa! Seluruh administrasi siswa telah LUNAS 100%. Tidak ada tunggakan.")

    # --- MODUL 7: MANAJEMEN AKUN GURU ---
    elif pilihan == "Manajemen Akun Guru" and st.session_state['role'] == 'admin':
        st.header("👥 Manajemen Akun Staf & Guru")
        tab_tambah, tab_hapus = st.tabs(["➕ Tambah Akun Baru", "❌ Hapus Akun"])
        with tab_tambah:
            with st.form("form_add_user", clear_on_submit=True):
                new_username = st.text_input("Username Baru (Tanpa Spasi)").strip().lower()
                new_password = st.text_input("Password Baru", type="password")
                new_role = st.selectbox("Hak Akses / Peran", ["guru", "admin"])
                if st.form_submit_button("Daftarkan Pengguna"):
                    if new_username and new_password:
                        try:
                            conn = get_connection()
                            cursor = conn.cursor()
                            hashed_pw = bcrypt.hashpw(new_password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')
                            cursor.execute("INSERT INTO users (username, password_hash, role) VALUES (?, ?, ?)", (new_username, hashed_pw, new_role))
                            conn.commit()
                            conn.close()
                            st.success(f"Akun {new_username} sukses dibuat!")
                            st.rerun()
                        except sqlite3.IntegrityError:
                            st.error("❌ Username sudah dipakai.")
                    else:
                        st.error("❌ Semua kolom wajib diisi.")
        with tab_hapus:
            conn = get_connection()
            df_list_users = pd.read_sql_query("SELECT username FROM users WHERE username != 'admin'", conn)
            conn.close()
            if df_list_users.empty:
                st.info("Tidak ada akun tambahan untuk dihapus.")
            else:
                with st.form("form_delete_user"):
                    user_to_delete = st.selectbox("Pilih Username yang akan Dihapus", df_list_users['username'])
                    if st.form_submit_button("🔴 Hapus Akun Permanen"):
                        if user_to_delete == st.session_state['username']:
                            st.error("❌ Tidak bisa menghapus akun sendiri!")
                        else:
                            conn = get_connection()
                            cursor = conn.cursor()
                            cursor.execute("DELETE FROM users WHERE username = ?", (user_to_delete,))
                            conn.commit()
                            conn.close()
                            st.success(f"Akun '{user_to_delete}' dihapus!")
                            st.rerun()
                            
        st.markdown("### 📋 Daftar Pengguna")
        conn = get_connection()
        df_users = pd.read_sql_query("SELECT id, username, role FROM users", conn)
        conn.close()
        st.dataframe(df_users, use_container_width=True)
        tampilkan_tombol_download(df_users, "Daftar_Pengguna", "DAFTAR PENGGUNA SISTEM")

    # --- MODUL 8: KOREKSI & HAPUS DATA ---
    elif pilihan == "Koreksi & Hapus Data" and st.session_state['role'] == 'admin':
        st.header("🛠️ Pusat Koreksi & Hapus Data")
        tab_del_absen, tab_del_nilai = st.tabs(["📝 Hapus Absensi Siswa", "📈 Hapus Nilai Siswa"])
        with tab_del_absen:
            conn = get_connection()
            df_absen = pd.read_sql_query("SELECT id, tanggal, nis, nama_siswa, kelas, status_kehadiran, petugas FROM absensi_siswa", conn)
            conn.close()
            if df_absen.empty:
                st.info("Belum ada data absensi.")
            else:
                st.dataframe(df_absen, use_container_width=True)
                tampilkan_tombol_download(df_absen, "DB_Absensi", "DATABASE ABSENSI")
                with st.form("form_hapus_absen"):
                    id_hapus_absen = st.number_input("Nomor ID Absensi yang dihapus", min_value=1, step=1)
                    if st.form_submit_button("🔴 Hapus Absen Permanen"):
                        conn = get_connection()
                        cursor = conn.cursor()
                        cursor.execute("SELECT id FROM absensi_siswa WHERE id = ?", (id_hapus_absen,))
                        if cursor.fetchone():
                            cursor.execute("DELETE FROM absensi_siswa WHERE id = ?", (id_hapus_absen,))
                            conn.commit()
                            st.success("Berhasil dihapus!")
                            st.rerun()
                        else: st.error("ID tidak ditemukan.")
                        conn.close()
                        
        with tab_del_nilai:
            conn = get_connection()
            df_nilai = pd.read_sql_query("SELECT id, tanggal_input, kelas, mata_pelajaran, nis, nama_siswa, nilai_akhir, guru_penginput FROM buku_nilai", conn)
            conn.close()
            if df_nilai.empty:
                st.info("Belum ada data nilai.")
            else:
                if 'nilai_akhir' in df_nilai.columns:
                    df_nilai['nilai_akhir'] = df_nilai['nilai_akhir'].round().astype(int)
                    
                st.dataframe(df_nilai, use_container_width=True)
                tampilkan_tombol_download(df_nilai, "DB_Nilai", "DATABASE NILAI")
                with st.form("form_hapus_nilai"):
                    id_hapus_nilai = st.number_input("Nomor ID Nilai yang dihapus", min_value=1, step=1)
                    if st.form_submit_button("🔴 Hapus Nilai Permanen"):
                        conn = get_connection()
                        cursor = conn.cursor()
                        cursor.execute("SELECT id FROM buku_nilai WHERE id = ?", (id_hapus_nilai,))
                        if cursor.fetchone():
                            cursor.execute("DELETE FROM buku_nilai WHERE id = ?", (id_hapus_nilai,))
                            conn.commit()
                            st.success("Berhasil dihapus!")
                            st.rerun()
                        else: st.error("ID tidak ditemukan.")
                        conn.close()

    # --- MODUL 9: BACKUP & RESTORE DATA ---
    elif pilihan == "Backup & Restore Data" and st.session_state['role'] == 'admin':
        st.header("🗄️ Backup & Pemulihan Database")
        col1, col2 = st.columns(2)
        with col1:
            st.subheader("1. Download Database")
            if os.path.exists(DB_FILE):
                with open(DB_FILE, "rb") as f:
                    st.download_button("📥 Download Database (.db)", f.read(), f"BACKUP_{datetime.now().strftime('%Y%m%d_%H%M%S')}.db", "application/x-sqlite3")
        with col2:
            st.subheader("2. Restore Database")
            uploaded_db = st.file_uploader("Unggah file (.db)", type=["db"])
            if uploaded_db and st.button("🔴 KONFIRMASI TIMPA DATA"):
                with open(DB_FILE, "wb") as f:
                    f.write(uploaded_db.getbuffer())
                st.success("Database berhasil dipulihkan! Me-restart sistem...")
                st.rerun()
